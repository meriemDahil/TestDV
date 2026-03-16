"""
reporters/word_report.py
------------------------
Reads validation_report.json and renders a comprehensive Word document
containing every metric, hash, stat, diff sample, and detail from the report.

Usage — standalone (no args needed, finds report automatically):
    python reporters/word_report.py
    python reporters/word_report.py --json validation_report.json
    python reporters/word_report.py --json report.json --out output.docx

Exit codes:  0 = success  |  1 = JSON not found  |  2 = generation failed
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# ─────────────────────────────────────────────
# Color palette
# ─────────────────────────────────────────────

class C:
    NAVY        = RGBColor(0x1B, 0x2A, 0x4A)
    TEAL        = RGBColor(0x0F, 0x6E, 0x56)
    TEAL_LIGHT  = RGBColor(0xE1, 0xF5, 0xEE)
    GREEN       = RGBColor(0x27, 0x50, 0x0A)
    GREEN_LIGHT = RGBColor(0xEA, 0xF3, 0xDE)
    RED         = RGBColor(0xA3, 0x2D, 0x2D)
    RED_LIGHT   = RGBColor(0xFC, 0xEB, 0xEB)
    AMBER       = RGBColor(0x85, 0x4F, 0x0B)
    AMBER_LIGHT = RGBColor(0xFA, 0xEE, 0xDA)
    GRAY        = RGBColor(0x5F, 0x5E, 0x5A)
    GRAY_LIGHT  = RGBColor(0xF1, 0xEF, 0xE8)
    GRAY_MID    = RGBColor(0xD3, 0xD1, 0xC7)
    WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
    BLACK       = RGBColor(0x00, 0x00, 0x00)
    MONO        = RGBColor(0x2C, 0x2C, 0x2A)


# ─────────────────────────────────────────────
# XML helpers
# ─────────────────────────────────────────────

def _rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _rgb_hex(rgb))
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _remove_borders(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcB.append(el)
    tcPr.append(tcB)


def _divider(doc: Document, color: str = "D3D1C7") -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ─────────────────────────────────────────────
# Paragraph / text helpers
# ─────────────────────────────────────────────

def _p(doc: Document, text: str, *,
       bold=False, size=10, color: RGBColor = C.BLACK,
       align=WD_ALIGN_PARAGRAPH.LEFT,
       before=0, after=4, italic=False, mono=False) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    run = para.add_run(text)
    run.bold         = bold
    run.italic       = italic
    run.font.size    = Pt(size)
    run.font.color.rgb = color
    run.font.name    = "Courier New" if mono else "Arial"
    return para


def _heading(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 15, 2: 12, 3: 11}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold           = True
    run.font.name      = "Arial"
    run.font.size      = Pt(sizes.get(level, 11))
    run.font.color.rgb = C.NAVY


def _cell_write(cell, text: str, *,
                bold=False, size=9, color: RGBColor = C.BLACK,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                italic=False, mono=False) -> None:
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold         = bold
    r.italic       = italic
    r.font.size    = Pt(size)
    r.font.color.rgb = color
    r.font.name    = "Courier New" if mono else "Arial"


def _cell_add_line(cell, text: str, *,
                   bold=False, size=8, color: RGBColor = C.GRAY,
                   italic=False, mono=False) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.bold         = bold
    r.italic       = italic
    r.font.size    = Pt(size)
    r.font.color.rgb = color
    r.font.name    = "Courier New" if mono else "Arial"


# ─────────────────────────────────────────────
# Status helpers
# ─────────────────────────────────────────────

def _status_colors(status: str) -> tuple[RGBColor, RGBColor]:
    return {
        "PASSED":  (C.GREEN_LIGHT, C.GREEN),
        "FAILED":  (C.RED_LIGHT,   C.RED),
        "SKIPPED": (C.AMBER_LIGHT, C.AMBER),
    }.get(status.upper(), (C.GRAY_LIGHT, C.GRAY))


def _fmt(value) -> str:
    """Format any scalar value for display."""
    if value is None:
        return "—"
    if isinstance(value, float):
        return f"{value:,.6g}"
    if isinstance(value, bool):
        return str(value)
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def _truncate(s: str, n: int = 60) -> str:
    s = str(s)
    return s if len(s) <= n else s[:n] + "…"


# ─────────────────────────────────────────────
# Reusable table builders
# ─────────────────────────────────────────────

def _kv_table(doc: Document, rows: list[tuple[str, str]],
              col_widths=(2.0, 4.6)) -> None:
    """Two-column key-value table."""
    if not rows:
        return
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        row = tbl.add_row()
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])
        fill = C.GRAY_LIGHT if i % 2 == 0 else C.WHITE
        _set_cell_bg(row.cells[0], fill)
        _set_cell_bg(row.cells[1], fill)
        _set_cell_margins(row.cells[0], top=60, bottom=60, left=120, right=80)
        _set_cell_margins(row.cells[1], top=60, bottom=60, left=120, right=80)
        _cell_write(row.cells[0], k, bold=True, size=9, color=C.NAVY)
        _cell_write(row.cells[1], v, size=9, color=C.BLACK, mono=True)
    doc.add_paragraph()


def _stat_table(doc: Document, src_stats: dict, tgt_stats: dict,
                issues: list, source: str, target: str) -> None:
    """Side-by-side stat comparison table."""
    stat_keys = ["min", "max", "mean", "stddev", "null_ratio", "non_null"] + \
                [k for k in src_stats if k.startswith("p")]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for ci, (lbl, w) in enumerate([("Metric", 1.2), (source, 1.8), (target, 1.8), ("Status", 0.8)]):
        c = tbl.rows[0].cells[ci]
        c.width = Inches(w)
        _set_cell_bg(c, C.GRAY_LIGHT)
        _set_cell_margins(c, top=60, bottom=60)
        _cell_write(c, lbl, bold=True, size=9, color=C.GRAY)
    issue_keys = {i["stat"] for i in issues}
    for key in stat_keys:
        sv = src_stats.get(key)
        tv = tgt_stats.get(key)
        if sv is None and tv is None:
            continue
        row = tbl.add_row()
        for ci, w in enumerate([1.2, 1.8, 1.8, 0.8]):
            row.cells[ci].width = Inches(w)
            _set_cell_margins(row.cells[ci], top=50, bottom=50, left=100, right=80)
        is_issue = key in issue_keys
        bg = C.RED_LIGHT if is_issue else C.WHITE
        for ci in range(4):
            _set_cell_bg(row.cells[ci], bg)
        _cell_write(row.cells[0], key, bold=True, size=9, color=C.NAVY)
        _cell_write(row.cells[1], _fmt(sv), size=9, mono=True)
        _cell_write(row.cells[2], _fmt(tv), size=9, mono=True)
        st = "FAIL" if is_issue else "OK"
        sc = C.RED if is_issue else C.GREEN
        _cell_write(row.cells[3], st, bold=True, size=9, color=sc,
                    align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


# ─────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────

def _build_cover(doc: Document, report: dict) -> None:
    summary = report.get("summary", {})
    status  = report.get("overall_status", "UNKNOWN")
    bg, fg  = _status_colors(status)

    # Banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, C.NAVY)
    _remove_borders(cell)
    _set_cell_margins(cell, top=200, bottom=200, left=240, right=240)
    _cell_write(cell,
                f"Validation Report  —  {report.get('source_label','?')}  →  {report.get('target_label','?')}",
                bold=True, size=14, color=C.WHITE)
    _cell_add_line(cell, f"Run ID: {report.get('run_id','—')}  ·  {report.get('timestamp','—')}",
                   size=9, color=RGBColor(0x9F, 0xE1, 0xCB))

    doc.add_paragraph()

    # Four metric cards
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    overall_bg, overall_fg = _status_colors(status)
    metrics = [
        ("Overall status", status,                             overall_bg, overall_fg),
        ("Pass rate",       f"{summary.get('pass_rate',0):.1%}", C.GRAY_LIGHT, C.NAVY),
        ("Checks",          f"{summary.get('passed',0)} passed / {summary.get('failed',0)} failed / {summary.get('skipped',0)} skipped",
                                                               C.GRAY_LIGHT, C.NAVY),
        ("Duration",        f"{report.get('total_duration_ms',0):.1f} ms", C.GRAY_LIGHT, C.NAVY),
    ]
    for i, (label, value, bg_c, fg_c) in enumerate(metrics):
        cell = tbl2.cell(0, i)
        cell.width = Inches(1.65)
        _set_cell_bg(cell, bg_c)
        _remove_borders(cell)
        _set_cell_margins(cell, top=120, bottom=120, left=120, right=80)
        cell.paragraphs[0].clear()
        r1 = cell.paragraphs[0].add_run(label)
        r1.font.name = "Arial"; r1.font.size = Pt(8); r1.font.color.rgb = C.GRAY
        p2 = cell.add_paragraph()
        r2 = p2.add_run(value)
        r2.bold = True; r2.font.name = "Arial"
        r2.font.size = Pt(11 if label == "Checks" else 13)
        r2.font.color.rgb = fg_c

    doc.add_paragraph()
    _divider(doc)


# ─────────────────────────────────────────────
# Layer 1 — Structural
# ─────────────────────────────────────────────

def _build_layer1(doc: Document, layer: dict) -> None:
    _heading(doc, "Layer 1 — Structural validation", level=1)
    _p(doc, "Schema compatibility check. Runs first — if this fails the engine stops.",
       size=9, color=C.GRAY, after=6)

    for check in layer.get("checks", []):
        name    = check["check_name"]
        status  = check["status"]
        msg     = check["message"]
        details = check.get("details", {})
        bg, fg  = _status_colors(status)
        dur     = check.get("duration_ms", 0)

        _heading(doc, name.replace("_", " ").title(), level=3)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        sc = tbl.cell(0, 0); vc = tbl.cell(0, 1)
        sc.width = Inches(1.2); vc.width = Inches(5.4)
        _set_cell_bg(sc, bg); _set_cell_bg(vc, C.WHITE)
        _set_cell_margins(sc, top=60, bottom=60)
        _set_cell_margins(vc, top=60, bottom=60, left=120, right=80)
        _cell_write(sc, status, bold=True, size=9, color=fg, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_write(vc, msg, size=9)
        _cell_add_line(vc, f"{dur:.2f} ms", size=8, color=C.GRAY, italic=True)

        # Structural details
        if details.get("mismatches"):
            doc.add_paragraph()
            _p(doc, "Type mismatches:", bold=True, size=9, color=C.NAVY, after=2)
            rows = [(col, f"{v['source_dtype']}  →  {v['target_dtype']}")
                    for col, v in details["mismatches"].items()]
            _kv_table(doc, rows)

        if details.get("null_issues"):
            _p(doc, "Null count differences:", bold=True, size=9, color=C.NAVY, after=2)
            rows = [(col, f"source={v['source_nulls']}  target={v['target_nulls']}")
                    for col, v in details["null_issues"].items()]
            _kv_table(doc, rows)

        if details.get("missing_in_target"):
            _p(doc, f"Missing in target: {details['missing_in_target']}", size=9, color=C.RED)
        if details.get("extra_in_target"):
            _p(doc, f"Extra in target: {details['extra_in_target']}", size=9, color=C.RED)

    doc.add_paragraph()


# ─────────────────────────────────────────────
# Layer 2 — Data level
# ─────────────────────────────────────────────

def _build_layer2(doc: Document, layer: dict, source: str, target: str) -> None:
    _heading(doc, "Layer 2 — Data-level validation", level=1)
    _p(doc, "Row counts, hashes, duplicates, bucketed checksums, and row-aligned diff.",
       size=9, color=C.GRAY, after=6)

    for check in layer.get("checks", []):
        name    = check["check_name"]
        status  = check["status"]
        msg     = check["message"]
        details = check.get("details", {})
        bg, fg  = _status_colors(status)
        dur     = check.get("duration_ms", 0)

        _heading(doc, name.replace("_", " ").title(), level=3)

        # Status + message row
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        sc = tbl.cell(0, 0); vc = tbl.cell(0, 1)
        sc.width = Inches(1.2); vc.width = Inches(5.4)
        _set_cell_bg(sc, bg); _set_cell_bg(vc, C.WHITE)
        _set_cell_margins(sc, top=60, bottom=60)
        _set_cell_margins(vc, top=60, bottom=60, left=120, right=80)
        _cell_write(sc, status, bold=True, size=9, color=fg, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_write(vc, msg, size=9)
        _cell_add_line(vc, f"{dur:.2f} ms", size=8, color=C.GRAY, italic=True)

        # ── row_count ──────────────────────────────────────────
        if name == "row_count" and details:
            doc.add_paragraph()
            _kv_table(doc, [
                (f"{source} rows",  _fmt(details.get("source_rows"))),
                (f"{target} rows",  _fmt(details.get("target_rows"))),
                ("Delta",           _fmt(details.get("delta"))),
            ])

        # ── full_dataset_hash ──────────────────────────────────
        elif name == "full_dataset_hash" and details:
            doc.add_paragraph()
            _kv_table(doc, [
                (f"{source} hash", _truncate(details.get("source_hash",""), 64)),
                (f"{target} hash", _truncate(details.get("target_hash",""), 64)),
            ])

        # ── column_level_hash ─────────────────────────────────
        elif name == "column_level_hash" and details:
            col_hashes = details.get("column_hashes", {})
            mismatched = details.get("mismatched_columns", [])
            if col_hashes:
                doc.add_paragraph()
                tbl2 = doc.add_table(rows=1, cols=3)
                tbl2.style = "Table Grid"
                for ci, (lbl, w) in enumerate([("Column", 1.8), ("Match", 0.8), ("Hashes", 4.0)]):
                    hc = tbl2.rows[0].cells[ci]
                    hc.width = Inches(w)
                    _set_cell_bg(hc, C.GRAY_LIGHT)
                    _set_cell_margins(hc, top=50, bottom=50)
                    _cell_write(hc, lbl, bold=True, size=9, color=C.GRAY)
                for col, info in col_hashes.items():
                    match = info.get("match", True)
                    row = tbl2.add_row()
                    row.cells[0].width = Inches(1.8)
                    row.cells[1].width = Inches(0.8)
                    row.cells[2].width = Inches(4.0)
                    for ci in range(3):
                        _set_cell_margins(row.cells[ci], top=50, bottom=50, left=100, right=80)
                        _set_cell_bg(row.cells[ci], C.RED_LIGHT if not match else C.WHITE)
                    _cell_write(row.cells[0], col, bold=(col in mismatched), size=9,
                                color=C.RED if not match else C.BLACK)
                    _cell_write(row.cells[1], "MATCH" if match else "DIFFER",
                                bold=True, size=9,
                                color=C.GREEN if match else C.RED,
                                align=WD_ALIGN_PARAGRAPH.CENTER)
                    sh = _truncate(info.get("source_hash",""), 28)
                    th = _truncate(info.get("target_hash",""), 28)
                    _cell_write(row.cells[2], f"src: {sh}", size=8, mono=True)
                    _cell_add_line(row.cells[2], f"tgt: {th}", size=8, mono=True,
                                   color=C.RED if not match else C.GRAY)
                doc.add_paragraph()

        # ── duplicate_rows ────────────────────────────────────
        elif name == "duplicate_rows" and details:
            doc.add_paragraph()
            _kv_table(doc, [
                (f"{source} duplicates", _fmt(details.get("source_duplicates"))),
                (f"{target} duplicates", _fmt(details.get("target_duplicates"))),
            ])

        # ── bucketed_checksum ─────────────────────────────────
        elif name == "bucketed_checksum" and details:
            buckets = details.get("mismatched_buckets", [])
            if buckets:
                doc.add_paragraph()
                _p(doc, f"{len(buckets)} mismatched bucket(s):", bold=True, size=9,
                   color=C.NAVY, after=2)
                tbl3 = doc.add_table(rows=1, cols=4)
                tbl3.style = "Table Grid"
                for ci, (lbl, w) in enumerate([
                    ("Bucket", 0.8), (f"{source} rows", 1.2),
                    (f"{target} rows", 1.2), ("Hash diff", 3.4)
                ]):
                    hc = tbl3.rows[0].cells[ci]
                    hc.width = Inches(w)
                    _set_cell_bg(hc, C.GRAY_LIGHT)
                    _set_cell_margins(hc, top=50, bottom=50)
                    _cell_write(hc, lbl, bold=True, size=9, color=C.GRAY)
                for b in buckets:
                    row = tbl3.add_row()
                    for ci, (val, w) in enumerate([
                        (str(b.get("bucket","")), 0.8),
                        (_fmt(b.get("source_rows","")), 1.2),
                        (_fmt(b.get("target_rows","")), 1.2),
                        (f"src: {_truncate(b.get('source_hash',''),18)}  tgt: {_truncate(b.get('target_hash',''),18)}", 3.4),
                    ]):
                        row.cells[ci].width = Inches(w)
                        _set_cell_margins(row.cells[ci], top=50, bottom=50, left=100, right=80)
                        _set_cell_bg(row.cells[ci], C.RED_LIGHT)
                        _cell_write(row.cells[ci], val, size=9, mono=(ci==3))
                doc.add_paragraph()

        # ── row_aligned_diff ──────────────────────────────────
        elif name == "row_aligned_diff" and status == "FAILED":
            diff = details.get("diff_sample", [])
            if diff:
                doc.add_paragraph()
                _p(doc, f"Differing rows (up to 20 shown):", bold=True,
                   size=9, color=C.NAVY, after=2)
                cols = list(diff[0].keys())
                tbl4 = doc.add_table(rows=1, cols=len(cols))
                tbl4.style = "Table Grid"
                col_w = round(6.6 / len(cols), 2)
                for ci, col in enumerate(cols):
                    hc = tbl4.rows[0].cells[ci]
                    hc.width = Inches(col_w)
                    _set_cell_bg(hc, C.NAVY)
                    _set_cell_margins(hc, top=50, bottom=50, left=80, right=80)
                    _cell_write(hc, _truncate(col, 16), bold=True, size=8, color=C.WHITE)
                for row_data in diff:
                    row = tbl4.add_row()
                    for ci, col in enumerate(cols):
                        row.cells[ci].width = Inches(col_w)
                        _set_cell_margins(row.cells[ci], top=40, bottom=40, left=80, right=80)
                        val = _truncate(_fmt(row_data.get(col)), 20)
                        is_diff_col = col.endswith("_src") and col[:-4]+"_tgt" in row_data
                        if is_diff_col:
                            src_v = row_data.get(col)
                            tgt_k = col[:-4] + "_tgt"
                            tgt_v = row_data.get(tgt_k)
                            is_diff = str(src_v) != str(tgt_v)
                            _set_cell_bg(row.cells[ci], C.AMBER_LIGHT if is_diff else C.WHITE)
                        _cell_write(row.cells[ci], val, size=8, mono=True)
                doc.add_paragraph()

        elif name == "row_aligned_diff" and status == "SKIPPED":
            doc.add_paragraph()
            _p(doc, f"To enable: {check.get('skipped_reason','')}", size=8,
               color=C.AMBER, italic=True, after=4)

    doc.add_paragraph()


# ─────────────────────────────────────────────
# Layer 3 — Business rules
# ─────────────────────────────────────────────

def _build_layer3(doc: Document, layer: dict, source: str, target: str) -> None:
    _heading(doc, "Layer 3 — Business rule validation", level=1)
    _p(doc, "Aggregation totals, float tolerance, group-by consistency, custom rules.",
       size=9, color=C.GRAY, after=6)

    for check in layer.get("checks", []):
        name    = check["check_name"]
        status  = check["status"]
        msg     = check["message"]
        details = check.get("details", {})
        bg, fg  = _status_colors(status)
        dur     = check.get("duration_ms", 0)

        _heading(doc, name.replace("_", " ").title(), level=3)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        sc = tbl.cell(0, 0); vc = tbl.cell(0, 1)
        sc.width = Inches(1.2); vc.width = Inches(5.4)
        _set_cell_bg(sc, bg); _set_cell_bg(vc, C.WHITE)
        _set_cell_margins(sc, top=60, bottom=60)
        _set_cell_margins(vc, top=60, bottom=60, left=120, right=80)
        _cell_write(sc, status, bold=True, size=9, color=fg, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_write(vc, msg, size=9)
        _cell_add_line(vc, f"{dur:.2f} ms", size=8, color=C.GRAY, italic=True)

        # ── aggregation issues ────────────────────────────────
        if name == "grand_total_aggregations":
            issues = details.get("issues", [])
            if issues:
                doc.add_paragraph()
                _p(doc, "Aggregation differences:", bold=True, size=9, color=C.NAVY, after=2)
                _kv_table(doc, [
                    (f"{i['column']} {i['function']}",
                     f"{source}={_fmt(i.get('source'))}  {target}={_fmt(i.get('target'))}  delta={_fmt(i.get('delta'))}")
                    for i in issues
                ])

        # ── float tolerance ───────────────────────────────────
        elif name == "float_tolerance":
            checked = details.get("columns_checked", [])
            issues  = details.get("issues", [])
            if checked:
                doc.add_paragraph()
                _p(doc, f"Columns checked: {', '.join(checked)}", size=9, color=C.GRAY, after=2)
            if issues:
                _p(doc, "Tolerance violations:", bold=True, size=9, color=C.NAVY, after=2)
                rows = []
                for i in issues:
                    rows.append((i["column"],
                                 f"max_abs={_fmt(i.get('max_absolute_diff'))}  "
                                 f"max_rel={_fmt(i.get('max_relative_diff'))}  "
                                 f"(tol abs={_fmt(i.get('abs_tolerance'))} "
                                 f"rel={_fmt(i.get('rel_tolerance'))})"))
                _kv_table(doc, rows)

        # ── group-by issues ───────────────────────────────────
        elif name == "group_by_consistency":
            issues = details.get("issues", [])
            if issues:
                doc.add_paragraph()
                for issue in issues:
                    if "group_count_mismatch" in str(issue):
                        _p(doc, f"Group count: {source}={issue.get('source_groups')}  "
                                f"{target}={issue.get('target_groups')}", size=9, color=C.RED)
                    elif "column" in issue:
                        grps = issue.get("mismatched_groups", [])
                        _p(doc, f"Column '{issue['column']}' — mismatched groups:",
                           bold=True, size=9, color=C.NAVY, after=2)
                        for g in grps[:10]:
                            _p(doc, f"  {g}", size=9, color=C.RED, mono=True, after=1)

        # ── custom rules ──────────────────────────────────────
        elif name.startswith("rule_"):
            if details:
                doc.add_paragraph()
                _kv_table(doc, [
                    (f"{source} value", _fmt(details.get("source"))),
                    (f"{target} value", _fmt(details.get("target"))),
                ])

        elif status == "SKIPPED":
            doc.add_paragraph()
            _p(doc, check.get("skipped_reason", ""), size=8,
               color=C.AMBER, italic=True, after=4)

    doc.add_paragraph()


# ─────────────────────────────────────────────
# Layer 4 — Statistical
# ─────────────────────────────────────────────

def _build_layer4(doc: Document, layer: dict, source: str, target: str) -> None:
    _heading(doc, "Layer 4 — Statistical validation", level=1)
    _p(doc, "Distribution profiles: min, max, mean, stddev, percentiles, null ratio, cardinality.",
       size=9, color=C.GRAY, after=6)

    for check in layer.get("checks", []):
        name    = check["check_name"]
        status  = check["status"]
        msg     = check["message"]
        details = check.get("details", {})
        bg, fg  = _status_colors(status)
        dur     = check.get("duration_ms", 0)

        _heading(doc, name.replace("_", " ").title(), level=3)

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        sc = tbl.cell(0, 0); vc = tbl.cell(0, 1)
        sc.width = Inches(1.2); vc.width = Inches(5.4)
        _set_cell_bg(sc, bg); _set_cell_bg(vc, C.WHITE)
        _set_cell_margins(sc, top=60, bottom=60)
        _set_cell_margins(vc, top=60, bottom=60, left=120, right=80)
        _cell_write(sc, status, bold=True, size=9, color=fg, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_write(vc, msg, size=9)
        _cell_add_line(vc, f"{dur:.2f} ms", size=8, color=C.GRAY, italic=True)

        # ── per-column stats ──────────────────────────────────
        if name == "per_column_statistics":
            profiles = details.get("column_profiles", {})
            col_issues = {i["column"] for i in details.get("issues", [])}

            for col, data in profiles.items():
                src_s = data.get("source", {})
                tgt_s = data.get("target", {})
                col_issue_list = data.get("issues", [])
                has_issue = col in col_issues

                doc.add_paragraph()
                _heading(doc,
                         f"  {col}" + ("  [DEVIATION]" if has_issue else "  [OK]"),
                         level=3)
                _stat_table(doc, src_s, tgt_s, col_issue_list, source, target)

                if col_issue_list:
                    _p(doc, "Deviations:", bold=True, size=9, color=C.RED, after=2)
                    for iss in col_issue_list:
                        _p(doc,
                           f"  {iss['stat']}: {source}={_fmt(iss['source'])}  "
                           f"{target}={_fmt(iss['target'])}  "
                           f"rel_diff={_fmt(iss.get('relative_diff'))}  "
                           f"(tol={_fmt(iss.get('tolerance'))})",
                           size=8, color=C.RED, mono=True, after=1)

        # ── null ratio ────────────────────────────────────────
        elif name == "null_ratio":
            issues = details.get("issues", [])
            if issues:
                doc.add_paragraph()
                _kv_table(doc, [
                    (i["column"],
                     f"{source}={_fmt(i['source_null_ratio'])}  "
                     f"{target}={_fmt(i['target_null_ratio'])}  "
                     f"delta={_fmt(i['delta'])}")
                    for i in issues
                ])

        # ── cardinality ───────────────────────────────────────
        elif name == "cardinality":
            issues = details.get("issues", [])
            if issues:
                doc.add_paragraph()
                _kv_table(doc, [
                    (i["column"],
                     f"{source}={_fmt(i['source_cardinality'])}  "
                     f"{target}={_fmt(i['target_cardinality'])}  "
                     f"delta={_fmt(i['delta'])}")
                    for i in issues
                ])

        # ── confidence score ──────────────────────────────────
        elif name == "distribution_confidence":
            if details:
                doc.add_paragraph()
                score = details.get("score", 0)
                _kv_table(doc, [
                    ("Score",     f"{score:.1%}"),
                    ("Threshold", f"{details.get('threshold', 0.95):.1%}"),
                    ("Result",    "ABOVE threshold" if score >= details.get("threshold", 0.95)
                                  else "BELOW threshold"),
                ])

    doc.add_paragraph()


# ─────────────────────────────────────────────
# Failures summary
# ─────────────────────────────────────────────

def _build_failures(doc: Document, failures: list) -> None:
    if not failures:
        return
    _heading(doc, "Failures summary", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for ci, (lbl, w) in enumerate([("Layer", 1.4), ("Check", 2.0), ("Message", 3.2)]):
        hc = tbl.rows[0].cells[ci]
        hc.width = Inches(w)
        _set_cell_bg(hc, C.RED_LIGHT)
        _set_cell_margins(hc, top=60, bottom=60)
        _cell_write(hc, lbl, bold=True, size=9, color=C.RED)
    for f in failures:
        row = tbl.add_row()
        for ci, (val, w) in enumerate([
            (f.get("layer",""),   1.4),
            (f.get("check",""),   2.0),
            (f.get("message",""), 3.2),
        ]):
            row.cells[ci].width = Inches(w)
            _set_cell_margins(row.cells[ci], top=60, bottom=60, left=100, right=80)
            _cell_write(row.cells[ci], val, size=9)
    doc.add_paragraph()


# ─────────────────────────────────────────────
# Skipped checks
# ─────────────────────────────────────────────

def _build_skipped(doc: Document, skipped: list) -> None:
    if not skipped:
        return
    _heading(doc, "Skipped checks", level=1)
    _p(doc, "These checks were not run due to missing configuration.", size=9, color=C.GRAY, after=6)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    for ci, (lbl, w) in enumerate([("Check", 2.2), ("How to enable", 4.4)]):
        hc = tbl.rows[0].cells[ci]
        hc.width = Inches(w)
        _set_cell_bg(hc, C.AMBER_LIGHT)
        _set_cell_margins(hc, top=60, bottom=60)
        _cell_write(hc, lbl, bold=True, size=9, color=C.AMBER)
    for sk in skipped:
        row = tbl.add_row()
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.4)
        _set_cell_margins(row.cells[0], top=60, bottom=60, left=120, right=80)
        _set_cell_margins(row.cells[1], top=60, bottom=60, left=120, right=80)
        _cell_write(row.cells[0], sk.get("check",""), bold=True, size=9, color=C.NAVY)
        _cell_write(row.cells[1], sk.get("to_enable", sk.get("message","")),
                    size=9, italic=True, color=C.GRAY)
    doc.add_paragraph()


# ─────────────────────────────────────────────
# AI narrative
# ─────────────────────────────────────────────

def _build_narrative(doc: Document, narrative: str) -> None:
    if not narrative or narrative.startswith("[AI narrative unavailable"):
        return
    _heading(doc, "AI narrative", level=1)

    # Banner
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, C.TEAL)
    _remove_borders(cell)
    _set_cell_margins(cell, top=100, bottom=100, left=200, right=200)
    _cell_write(cell,
                "Generated by Azure OpenAI from the completed report — "
                "the AI did not perform any validation logic.",
                bold=True, size=9, color=C.WHITE)
    doc.add_paragraph()

    for line in narrative.splitlines():
        if line.strip():
            _p(doc, line, size=10, after=4)
        else:
            doc.add_paragraph()


# ─────────────────────────────────────────────
# Main entry point
# ─────────────────────────────────────────────

def generate_word_report(
    json_path: str | Path,
    output_path: str | Path | None = None,
) -> Path:
    json_path = Path(json_path)
    if output_path is None:
        output_path = json_path.with_suffix(".docx")
    output_path = Path(output_path)

    with open(json_path, encoding="utf-8") as f:
        report = json.load(f)

    summary  = report.get("summary", {})
    layers   = report.get("layers",  [])
    failures = summary.get("failures", [])
    skipped  = summary.get("skipped_checks", [])
    source   = report.get("source_label", "Source")
    target   = report.get("target_label", "Target")

    # Layer lookup by name prefix
    def get_layer(prefix: str) -> dict:
        return next((l for l in layers if l["layer_name"].startswith(prefix)), {})

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    _build_cover(doc, report)

    l1 = get_layer("1_")
    if l1:
        _build_layer1(doc, l1)
        _divider(doc)

    l2 = get_layer("2_")
    if l2:
        _build_layer2(doc, l2, source, target)
        _divider(doc)

    l3 = get_layer("3_")
    if l3:
        _build_layer3(doc, l3, source, target)
        _divider(doc)

    l4 = get_layer("4_")
    if l4:
        _build_layer4(doc, l4, source, target)
        _divider(doc)

    _build_failures(doc, failures)
    _build_skipped(doc, skipped)
    _build_narrative(doc, report.get("ai_narrative", ""))

    doc.save(output_path)
    print(f"[word_report] Written → {output_path}")
    return output_path


# ─────────────────────────────────────────────
# Standalone CLI
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    _DEFAULT_JSON = Path(__file__).resolve().parent.parent / "validation_report.json"

    ap = argparse.ArgumentParser(
        description="Generate a comprehensive Word report from a validation_report.json.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
examples:
  python reporters/word_report.py                                  # uses default: {_DEFAULT_JSON}
  python reporters/word_report.py --json validation_report.json
  python reporters/word_report.py --json /path/to/report.json --out /path/to/output.docx
        """,
    )
    ap.add_argument("--json", default=str(_DEFAULT_JSON), metavar="PATH",
                    help=f"Path to validation_report.json. Defaults to {_DEFAULT_JSON}")
    ap.add_argument("--out",  default=None, metavar="PATH",
                    help="Output .docx path. Defaults to same folder as --json with .docx extension.")
    args = ap.parse_args()

    json_path = Path(args.json)
    if not json_path.exists():
        print(f"[word_report] ERROR: JSON file not found: {json_path}", file=sys.stderr)
        print(f"[word_report] Run the validation engine first, or pass --json PATH", file=sys.stderr)
        sys.exit(1)

    try:
        out = generate_word_report(json_path, args.out)
        print(f"[word_report] Done — {out}")
        sys.exit(0)
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"[word_report] ERROR: {e}", file=sys.stderr)
        sys.exit(2)